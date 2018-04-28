#!/usr/bin/env python
# -*- coding: utf-8 -*-
'''
Created on 2016年11月20日
工具模块
@author: dell
'''
import jsonpickle
from model import model
import time
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import zipfile,os
import logging
from config_default import logging_level
type = dict(choice="选择题", coding="编程题", filla="读程序写结果", fillb="程序填空", judge="判断")
in_state = ["未参加", "正在考试", "已结束", "作弊"]

def getFileRotatingLog():
    logger = None
    logger = logging.getLogger()
    # logger.setLevel(logging.DEBUG)
    logger.setLevel(logging_level)
    rh = logging.FileHandler("/home/exam/log/information.log")
    formatStr = "[%(asctime)s] %(filename)s:%(lineno)d(%(funcName)s): [%(levelname)s] %(message)s"
    format = logging.Formatter(formatStr)
    rh.setFormatter(format)
    logger.addHandler(rh)
    return logger

def dfs_get_zip_file(input_path,result):
    files = os.listdir(input_path)
    for file in files:
        if os.path.isdir(input_path+'/'+file):
            dfs_get_zip_file(input_path+'/'+file,result)
        else:
            result.append(input_path+'/'+file)

def zip_path(input_path,output_path,output_name):
    f = zipfile.ZipFile(output_path+'/'+output_name,'w',zipfile.ZIP_DEFLATED)
    filelists = []
    dfs_get_zip_file(input_path,filelists)
    for file in filelists:
        f.write(file)
    f.close()
    # print output_path+r"/"+output_name


def word(student_st_id,exam_ex_id,filepath):
    # 打开文档
    document = Document()
    # 加入不同等级的标题
    exam = model.Exam_model.getByPK(exam_ex_id)
    document.add_heading(u'\t\t' + exam['ex_name'], 0)

    information = model.Information_model()
    information_data = information.query('SELECT * FROM information WHERE \
                student_st_id=%s and exam_ex_id=%s' % (student_st_id, exam_ex_id))
    information = model.Information_model(**information_data[0])
    exam_question = model.Exam_question_model.getByArgs(information_in_id=information.in_id)
    choice_question = []
    judge_question = []
    # filla是读程序写结果
    filla_question = []
    fillb_question = []
    coding_question = []
    for item in exam_question:
        if item.eq_qt_type == 'choice':
            question_data = model.Question_model.getByPK(item.qt_id)
            choice_data = model.Choice_model.getByPK(item.qt_id)
            item = dict(item, **choice_data)
            item = dict(item, **question_data)
            choice_question.append(item)
        elif item.eq_qt_type == 'judge':
            question_data = model.Question_model.getByPK(item.qt_id)
            judge_data = model.Judge_model.getByPK(item.qt_id)
            item = dict(item, **judge_data)
            item = dict(item, **question_data)
            judge_question.append(item)
        elif item.eq_qt_type == 'filla':
            question_data = model.Question_model.getByPK(item.qt_id)
            filla_data = model.Filla_model.getByPK(item.qt_id)
            item = dict(item, **filla_data)
            item = dict(item, **question_data)
            filla_question.append(item)
        elif item.eq_qt_type == 'coding':
            question_data = model.Question_model.getByPK(item.qt_id)
            coding_data = model.Coding_model.getByPK(item.qt_id)
            item = dict(item, **coding_data)
            item = dict(item, **question_data)
            coding_question.append(item)
    fillb_qt = model.Exam_question_model.query('select distinct(qt_id),\
                                                            eq_pre_score from exam_question where information_in_id = %s \
                                                            and eq_qt_type = %s' % (
    information.in_id, "'" + 'fillb' + "'"))
    for item in fillb_qt:
        question_data = model.Question_model.getByPK(item['qt_id'])
        fillb_data = model.Fillb_model.getByPK(item['qt_id'])

        eq_id_data = model.Exam_question_model.query('select eq_id,eq_answer ,eq_get_score from \
                                                                        exam_question where qt_id = %s and information_in_id = %s' \
                                                     % (item['qt_id'], information.in_id))
        eq_id_data = [model.Exam_question_model(**items) for items in eq_id_data]
        i = 1
        fillb_coding = fillb_data.fb_pre_coding.split('&&&')
        print len(fillb_coding)
        for k in eq_id_data:
            j = 2 * i - 1
            fillb_coding[j] = u'空 ' + str(i)
            i += 1
        fillb_data.fb_pre_coding = ''.join(fillb_coding)

        item = dict(item, **fillb_data)
        item = dict(item, **question_data)
        fillb = []
        fillb.append(item)
        fillb.append(eq_id_data)
        # print fillb
        fillb_question.append(fillb)

    student = model.Student_model.getByPK(student_st_id)
    # 添加文本
    paragraph = document.add_paragraph()
    paragraph.add_run(u'学号:\t')
    paragraph.add_run(u'%s' % student['st_id'])
    paragraph.add_run(u'\t姓名:    ')
    paragraph.add_run(student['st_name'])
    paragraph.add_run(u'\t班级:    ')
    paragraph.add_run(student['st_specialty'])
    paragraph.add_run(u'\t得分:    ')
    if information['in_score'] < 0:
        # information['in_score']=u'未出分'
        paragraph.add_run(u'未出分')
    else:
        paragraph.add_run(str(information['in_score']))
    # #设置字号
    # run = paragraph.add_run(u'设置字号、')
    # run.font.size = Pt(24)
    #
    # #设置字体
    # run = paragraph.add_run('Set Font,')
    # run.font.name = 'Consolas'

    # 设置中文字体
    # run = paragraph.add_run('学号: '+student['st_id'])
    # run.font.name=u'宋体'
    # r = run._element
    # r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # #增加无序列表
    # document.add_paragraph(
    #     u'无序列表元素1', style='List Bullet'
    # )
    # document.add_paragraph(
    #     u'无序列表元素2', style='List Bullet'
    # )
    document.add_heading(u'选择题', 1)
    for item in choice_question:
        # print item
        if item['eq_get_score'] < 0:
            # item['eq_get_score'] = "未出分"
            document.add_paragraph(
                item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + u'未出分',
                style='List Number')
        else:
            document.add_paragraph(
                item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + str(item['eq_get_score']),
                style='List Number')

        document.add_paragraph(
            u'A. ' + item['cc_a'], style='List Bullet'
        )
        document.add_paragraph(
            u'B. ' + item['cc_b'], style='List Bullet'
        )
        document.add_paragraph(
            u'C. ' + item['cc_c'], style='List Bullet'
        )
        document.add_paragraph(
            u'D. ' + item['cc_d'], style='List Bullet'
        )

    document.add_heading(u'判断题', 1)
    for item in judge_question:
        # print item
        if item['eq_get_score'] < 0:
            # item['eq_get_score'] = "未出分"
            document.add_paragraph(
                item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + u'未出分',
                style='List Number')
        else:
            document.add_paragraph(
                item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + str(item['eq_get_score']),
                style='List Number')

    document.add_heading(u'读程序写结果', 1)
    for item in filla_question:
        # print item
        if item['eq_get_score'] < 0:
            # item['eq_get_score'] = "未出分"
            document.add_paragraph(
                item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + u'未出分',
                style='List Number')
        else:
            document.add_paragraph(
                item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + str(item['eq_get_score']),
                style='List Number')

    document.add_heading(u'程序填空', 1)
    for item in fillb_question:
        print item
        answer = ''
        i = 1
        for eq in item[1]:
            if eq['eq_get_score'] < 0:
                # eq['eq_get_score'] = u"未出分"
                answer += u'空' + str(i) + '\t' + eq.eq_answer + u'\t 得分: ' + u"未出分" + '\n'
            else:
                answer += u'空' + str(i) + '\t' + eq.eq_answer + u'\t 得分: ' + str(eq['eq_get_score']) + '\n'
            i += 1
        document.add_paragraph(
            item[0]['qt_stem'] + u'\n' + item[0]['fb_pre_coding'] + u'\n' + answer, style='List Number'
        )

    document.add_heading(u'编程题', 1)
    for item in coding_question:
        # print item
        if item['eq_get_score'] < 0:
            # item['eq_get_score'] = "未出分"
            document.add_paragraph(
                item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + u'未出分',
                style='List Number')
        else:
            document.add_paragraph(
                item['qt_stem'] + u'\n 学生答案: ' + item['eq_answer'] + u'\t 得分: ' + str(item['eq_get_score']),
                style='List Number')

    # #增加有序列表
    # document.add_paragraph(
    #     u'有序列表元素1', style='List Number'
    # )



    # 增加分页
    # document.add_page_break()

    # 保存文件
    # print "savepage"
    # document.save(u'../examTransplant2.2/source/exampage/2014.docx')
    # document.save(str(filepath))
    document.save(filepath)

def QuestionWord(question_list,question_back_id,filepath):
    print question_list
    # 打开文档
    document = Document()
    # 加入不同等级的标题
    questionback = model.Questions_bank_model.getByPK(question_back_id)
    document.add_heading(u'\t\t' + questionback.qb_name, 0)
    choice_question = []
    judge_question = []
    # filla是读程序写结果
    filla_question = []
    fillb_question = []
    coding_question = []
    for item in question_list:
        print item.qt_type.decode('utf-8')
        if str(item.qt_type) == 'choice':
            question_data = model.Question_model.getByPK(item.qt_id)
            choice_data = model.Choice_model.getByPK(item.qt_id)
            item = dict(item, **choice_data)
            item = dict(item, **question_data)
            choice_question.append(item)
        elif item.qt_type == 'judge':
            question_data = model.Question_model.getByPK(item.qt_id)
            judge_data = model.Judge_model.getByPK(item.qt_id)
            item = dict(item, **judge_data)
            item = dict(item, **question_data)
            judge_question.append(item)
        elif item.qt_type == 'filla':
            question_data = model.Question_model.getByPK(item.qt_id)
            filla_data = model.Filla_model.getByPK(item.qt_id)
            item = dict(item, **filla_data)
            item = dict(item, **question_data)
            filla_question.append(item)
        elif item.qt_type == 'coding':
            question_data = model.Question_model.getByPK(item.qt_id)
            coding_data = model.Coding_model.getByPK(item.qt_id)
            item = dict(item, **coding_data)
            item = dict(item, **question_data)
            coding_question.append(item)
        elif item.qt_type == 'fillb':
            question_data = model.Question_model.getByPK(item.qt_id)
            fillb_data = model.Fillb_model.getByPK(item.qt_id)
            fillb_coding = fillb_data.fb_pre_coding.split('&&&')
            print len(fillb_coding)
            j=1
            for i in range(1,len(fillb_coding),2):
                fillb_coding[i] = u'空 ' + str(j)
                j+=1
            fillb_data.fb_pre_coding = ''.join(fillb_coding)
            item = dict(item, **fillb_data)
            item = dict(item, **question_data)
            fillb_question.append(item)
    document.add_heading(u'选择题', 1)
    for item in choice_question:
        document.add_paragraph(
                item['qt_stem']+ u'\n 答案: ' + item['cc_answer'] + u'\n 准确率: ' + item['qt_pre_rate'],
                style='List Number')

        document.add_paragraph(
            u'A. ' + item['cc_a'], style='List Bullet'
        )
        document.add_paragraph(
            u'B. ' + item['cc_b'], style='List Bullet'
        )
        document.add_paragraph(
            u'C. ' + item['cc_c'], style='List Bullet'
        )
        document.add_paragraph(
            u'D. ' + item['cc_d'], style='List Bullet'
        )

    document.add_heading(u'判断题', 1)
    for item in judge_question:
        document.add_paragraph(
                item['qt_stem']+ u'\n 答案: ' + item['jd_answer'] + u'\n 准确率: ' + item['qt_pre_rate'],
                style='List Number')

    document.add_heading(u'读程序写结果', 1)
    for item in filla_question:
        document.add_paragraph(
                item['qt_stem'] + u'\n 答案: ' + item['filla_answer'] + u'\n 准确率: ' + item['qt_pre_rate'],
                style='List Number')

    document.add_heading(u'程序填空', 1)
    for item in fillb_question:
        print item
        document.add_paragraph(
            item['qt_stem'] + u'\n' + item['fb_pre_coding'] + u'\n' + u'\n 准确率: ' + item['qt_pre_rate'], style='List Number'
        )

    document.add_heading(u'编程题', 1)
    for item in coding_question:
        document.add_paragraph(
                item['qt_stem'] + u'\n 预留代码: ' + item['co_test_coding'] + u'\n 准确率: ' + item['qt_pre_rate'],
                style='List Number')

    # #增加有序列表
    # document.add_paragraph(
    #     u'有序列表元素1', style='List Number'
    # )



    # 增加分页
    # document.add_page_break()

    # 保存文件
    # print "savepage"
    # document.save(u'../examTransplant2.2/source/exampage/2014.docx')
    document.save(str(filepath))
'''
对象转json
'''
def objtojson(obj):
    return jsonpickle.encode(obj)


def upInformationScore(in_id):
    exam_question = model.Exam_question_model.getByArgs(information_in_id=in_id)
    score = 0
    flag = 0
    for question in exam_question:
        if question.eq_get_score >= 0:
            score += question.eq_get_score
        else:
            flag = 1
            break
    if exam_question == None:
        score = -1
    if flag == 0:
        return score
    else:
        return -1


def GetScore(delay, in_id):
    while 1:
        time.sleep(delay)
        print "getscore"
        exam_question = model.Exam_question_model.getByArgs(information_in_id=in_id)
        score = 0
        flag = 0
        for question in exam_question:
            if question.eq_get_score >= 0:
                score += question.eq_get_score
            else:
                flag = 1
                break
        if flag == 0:
            information = model.Information_model()
            information.in_id = in_id
            information.in_score = score
            information.update()
            break
def SaveFillb(information_in_id):
    examquestion = model.Exam_question_model.query('select distinct(qt_id) from exam_question where information_in_id = %s and  eq_qt_type =%s'% \
                                                   (information_in_id,"'"+'fillb'+"'"))
    examquestion = [model.Exam_question_model(**item) for item in examquestion]
    if examquestion!=None:
        for fillb_qt in examquestion:
            fillb = model.Fillb_model.getByPK(fillb_qt['qt_id'])
            fillb_question = model.Exam_question_model.query('select * from exam_question where information_in_id = %s and  qt_id =%s'% \
                                                   (information_in_id,fillb_qt['qt_id']))
            fillb_question = [model.Exam_question_model(**item) for item in fillb_question]
            i = 1
            for item in fillb_question:
                j = 2 * i - 1
                fillb_coding = fillb.fb_pre_coding.split('&&&')
                fillb_coding[j] = item.eq_answer
                item.fillb_coding = ''.join(fillb_coding)
                item.update()
                i +=1







'''
must_params表示接口必须要有的参数，args代表页面传过来的参数可直接传web.input()
'''


def paramsok(must_params, args):
    for k in args:
        if k not in args:
            return Status.__params_not_ok__
    return Status.__params_ok__


'''
md5散列
'''


def md5(str):
    import hashlib
    m = hashlib.md5()
    m.update(str)
    return m.hexdigest()


'''
对接口状态进行定义
'''


class Status:
    __success__ = 1  # 操作成功
    __error__ = 0  # 操作失败
    __params_not_ok__ = 2  # 参数传递错误
    __params_ok__ = 3  # 参数传递正确
    __not_login__ = 4  # 还未登陆
    __password_not_match__ = 5  # 密码错误
    __obj_null__ = 6  # 对象为空
    __system_exception = 7  # 程序抛出异常
    __not_exist__ = 8  # 不存在


'''
对不需要使用分页的接口返回数据进行封装
status:接口返回状态,message:接口状态描述,body:返回数据
'''


class Response:
    def __init__(self, status=Status.__error__, message="unknow", body=None):
        self.status = status
        self.body = body
        self.message = message


'''
data:要填充的数据，如{"id":1,"name":"ajj"}
totalRow:总条数，
currentPage:当前请求页，pageSize:每页显示数目，
status:返回状态，在Status类中定义了各种状态类型
message:状态信息描述
'''


class Page:
    def __init__(self, data, totalRow, currentPage, pageSize=10, status=Status.__error__, message="未知"):
        self.totalRow = totalRow
        self.data = data
        self.currentPage = int(currentPage)
        self.pageSize = pageSize
        self.totalPage = totalRow / pageSize if (totalRow % pageSize == 0) else (totalRow / pageSize + 1)  # 总页数
        self.firstPage = self.currentPage == 1  # 是否为第一页
        self.lastPage = self.currentPage == self.totalPage  # 是否为最后一页
        self.status = status
        self.message = message


'''
test和write两个函数配合自动生成model类，建立与数据库之间的联系
'''


def test(db, file_name):
    file = open(file_name, 'w')
    sql_tables = 'show tables'
    sql_columns = 'select * from information_schema.COLUMNS where table_name = "%s"'
    sql_pri_column = 'select COLUMN_KEY,COLUMN_NAME from INFORMATION_SCHEMA.COLUMNS where table_name="%s" AND COLUMN_KEY="PRI"'
    sql_not_null_column = 'select * from INFORMATION_SCHEMA.COLUMNS where table_name= "%s" and IS_NULLABLE = "NO"'
    tables = db.query(sql_tables)
    for table in tables:
        __table__ = "'" + str(table['Tables_in_examdb']) + "'"
        table1 = table['Tables_in_examdb']
        __pk__ = db.query(
            'select COLUMN_KEY,COLUMN_NAME from INFORMATION_SCHEMA.COLUMNS where table_name=%s AND COLUMN_KEY="PRI"' % __table__)[
            0]['COLUMN_NAME']
        __notnull__ = "set({"
        columns = db.query('select * from information_schema.COLUMNS where table_name = %s' % __table__)
        __attrnum__ = len(columns)
        print __attrnum__
        __attr__ = "set(["
        __updateable__ = __insertable__ = "set({"
        __countperpage__ = 10
        for column in columns:
            column_name = column['COLUMN_NAME']
            __attr__ += "'" + column_name + "',"

            __insertable__ += ("'" + column_name + "',") if column_name != __pk__ else ""
            __updateable__ += ("'" + column_name + "',") if column_name != __pk__ else ""
            __notnull__ += ("'" + column_name + "',") if column['IS_NULLABLE'] == 'NO' and column_name != __pk__ else ""
        __attr__ += "])"
        __insertable__ += "})"
        __updateable__ += "})"
        __notnull__ += "})"

        write(file, table1, __pk__, __attr__, __insertable__, __updateable__, __notnull__)
        file.write('\n')
    file.close()


'''
配合上述test函数
'''


def write(file, __table__, __pk__, __attr__, __insertable__, __updateable__, __notnull__):
    file.write("class " + __table__[0].upper() + __table__[1:] + '_model(Model):\n')
    file.write("    __table__ = '" + __table__ + "'\n")
    file.write("    __pk__ = '" + __pk__ + "'\n")
    file.write('    __attr__ = ' + __attr__ + "\n")
    file.write('    __insertable__ = ' + __insertable__ + "\n")
    file.write('    __updateable__ = ' + __updateable__ + "\n")
    file.write('    __notnull__ = ' + __notnull__ + "\n")
    file.write('    __attrnum__ = len(__attr__)' + "\n")
    file.write('    __countperpage__ = 10' + "\n")
